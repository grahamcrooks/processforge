"""
Stage 4 — Process Analysis Summary Word document generator.
Produces a stakeholder-ready Process Analysis Summary (.docx) using python-docx.
Supports optional branding: org name, primary colour, logo.
"""
import io
import base64
from datetime import date
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Colour helpers
# ---------------------------------------------------------------------------

DEFAULT_PRIMARY = "0057B8"   # Bupa blue


def _hex_to_rgb(hex_colour: str) -> RGBColor:
    h = hex_colour.lstrip("#")
    if len(h) != 6:
        h = DEFAULT_PRIMARY
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    return RGBColor(r, g, b)


def _set_cell_background(cell, hex_colour: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour.lstrip("#"))
    tcPr.append(shd)


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _add_heading(doc: Document, text: str, level: int, colour: RGBColor):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = colour
    return p


def _add_table_row(table, cells: list[str], bold: bool = False, bg_hex: Optional[str] = None):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = bold
        if bg_hex:
            _set_cell_background(cell, bg_hex)
    return row


# ---------------------------------------------------------------------------
# Main generator
# ---------------------------------------------------------------------------

def generate_bpin(process_data: dict, branding) -> bytes:
    org_name = (branding.org_name if branding else None) or "Bupa"
    colour_hex = (branding.primary_colour if branding else None) or DEFAULT_PRIMARY
    logo_b64 = branding.logo_base64 if branding else None
    primary = _hex_to_rgb(colour_hex)
    today = date.today().strftime("%-d %B %Y")

    process_name = process_data.get("process_name", "Process")
    lanes = process_data.get("lanes", [])
    steps = process_data.get("steps", [])
    entities = process_data.get("data_entities", [])
    integrations = process_data.get("integrations", [])
    decisions = process_data.get("decisions", [])

    doc = Document()

    # -----------------------------------------------------------------------
    # Page margins
    # -----------------------------------------------------------------------
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # -----------------------------------------------------------------------
    # Cover page
    # -----------------------------------------------------------------------
    # Logo (if provided)
    if logo_b64:
        try:
            logo_bytes = base64.b64decode(logo_b64)
            logo_stream = io.BytesIO(logo_bytes)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run()
            run.add_picture(logo_stream, width=Cm(4))
        except Exception:
            pass  # Invalid logo — skip silently

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(f"Process Analysis Summary")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = primary

    doc.add_paragraph()

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run(process_name)
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(18)

    doc.add_paragraph()

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.add_run(f"{org_name}\n{today}").font.size = Pt(12)

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # 1. Executive Summary
    # -----------------------------------------------------------------------
    _add_heading(doc, "1. Executive Summary", 1, primary)
    doc.add_paragraph(
        f"This Process Analysis Summary describes the technical artefacts "
        f"generated for the {process_name} process at {org_name}. "
        f"The artefacts were produced by analysing process flow diagrams using AI-assisted "
        f"extraction and are intended to accelerate implementation."
    )

    # -----------------------------------------------------------------------
    # 2. Process Overview
    # -----------------------------------------------------------------------
    _add_heading(doc, "2. Process Overview", 1, primary)

    _add_heading(doc, "2.1 Process Name", 2, primary)
    doc.add_paragraph(process_name)

    _add_heading(doc, "2.2 Swimlanes / Roles", 2, primary)
    if lanes:
        for lane in lanes:
            doc.add_paragraph(lane, style="List Bullet")
    else:
        doc.add_paragraph("No swimlanes identified.")

    # -----------------------------------------------------------------------
    # 3. Process Steps
    # -----------------------------------------------------------------------
    _add_heading(doc, "3. Process Steps", 1, primary)

    if steps:
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = "Table Grid"

        # Header row
        hdr = table.rows[0]
        headers = ["ID", "Type", "Name", "Lane"]
        for i, h in enumerate(headers):
            cell = hdr.cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            _set_cell_background(cell, colour_hex)

        for step in steps:
            _add_table_row(table, [
                step.get("id", ""),
                step.get("type", ""),
                step.get("name", ""),
                step.get("lane", ""),
            ])
    else:
        doc.add_paragraph("No steps identified.")

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # 4. Decision Points
    # -----------------------------------------------------------------------
    _add_heading(doc, "4. Decision Points", 1, primary)

    if decisions:
        for d in decisions:
            _add_heading(doc, d.get("question", "Decision"), 2, primary)
            outcomes = d.get("outcomes", [])
            if outcomes:
                for o in outcomes:
                    doc.add_paragraph(o, style="List Bullet")
            else:
                doc.add_paragraph("No outcomes recorded.")
    else:
        doc.add_paragraph("No decision points identified.")

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # 5. Data Entities
    # -----------------------------------------------------------------------
    _add_heading(doc, "5. Data Entities", 1, primary)

    if entities:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0]
        for i, h in enumerate(["Entity", "Notes"]):
            cell = hdr.cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            _set_cell_background(cell, colour_hex)
        for entity in entities:
            _add_table_row(table, [entity, ""])
    else:
        doc.add_paragraph("No data entities identified.")

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # 6. Integrations
    # -----------------------------------------------------------------------
    _add_heading(doc, "6. Integrations", 1, primary)

    if integrations:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0]
        for i, h in enumerate(["System", "Type", "Notes"]):
            cell = hdr.cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            _set_cell_background(cell, colour_hex)
        for intg in integrations:
            _add_table_row(table, [intg, "REST / SOAP", ""])
    else:
        doc.add_paragraph("No external integrations identified.")

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # 7. Generated Artefacts
    # -----------------------------------------------------------------------
    _add_heading(doc, "7. Generated Artefacts", 1, primary)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(["Artefact", "File", "Purpose"]):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_background(cell, colour_hex)

    import re as _re
    _slug = _re.sub(r"[^a-z0-9]+", "-", process_data.get("process_name", "process").lower()).strip("-")[:60]
    artefacts = [
        ("BPMN Process Model", f"{_slug}.bpmn", "Import into Blueprint as process definition"),
        ("PostgreSQL DDL", f"{_slug}-schema.sql", "Database schema for Blueprint data layer"),
        ("OpenAPI Specification", f"{_slug}-api-spec.yaml", "API contract for integration services"),
        ("Process Analysis Summary", "process-analysis-summary.docx", "This document — stakeholder reference"),
    ]
    for row_data in artefacts:
        _add_table_row(table, list(row_data))

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # 8. Next Steps
    # -----------------------------------------------------------------------
    _add_heading(doc, "8. Next Steps", 1, primary)
    next_steps = [
        "Review generated BPMN and validate against source diagrams.",
        "Confirm data entities and extend DDL with business-specific columns.",
        "Review OpenAPI spec with integration teams.",
        "Import BPMN into Blueprint and run initial configuration review.",
        "Schedule stakeholder walkthrough using this document.",
    ]
    for step in next_steps:
        doc.add_paragraph(step, style="List Number")

    # -----------------------------------------------------------------------
    # Footer note
    # -----------------------------------------------------------------------
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_run = footer_p.add_run(
        f"Generated by Process Forge · {today} · AI-assisted — review all content before use."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)

    # -----------------------------------------------------------------------
    # Serialise to bytes
    # -----------------------------------------------------------------------
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
